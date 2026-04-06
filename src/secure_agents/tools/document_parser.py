"""Document parser tool - extracts text from PDF and DOCX files securely.

When sandbox mode is enabled (the default), parsing runs inside a Docker
container with no network access. When disabled, parsing runs directly on
the host but with file validation (type, size, magic bytes).
"""

from __future__ import annotations

import json
from pathlib import Path
from typing import Any

import structlog

from secure_agents.core.base_tool import BaseTool
from secure_agents.core.registry import register_tool
from secure_agents.core.security import validate_file

logger = structlog.get_logger()

# Script that runs inside the Docker sandbox for PDF parsing
_SANDBOX_PARSE_SCRIPT = """
import json
import sys
from pathlib import Path

input_path = Path("/input/data.json")
data = json.loads(input_path.read_text())
file_path = Path("/input") / data["filename"]
file_type = data["file_type"]

result = {"text": "", "metadata": {}, "page_count": 0, "file_type": file_type}

try:
    if file_type == "pdf":
        import pdfplumber
        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            result["page_count"] = len(pdf.pages)
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        result["text"] = "\\n\\n".join(text_parts)
    elif file_type in ("docx", "doc"):
        import docx
        doc = docx.Document(str(file_path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        result["text"] = "\\n\\n".join(paragraphs)
        result["page_count"] = max(1, len(result["text"]) // 3000)
    result["metadata"] = {"filename": data["filename"], "size_bytes": file_path.stat().st_size}
except Exception as e:
    result["error"] = str(e)

output_path = Path("/output/result.json")
output_path.write_text(json.dumps(result))
"""


@register_tool("document_parser")
class DocumentParserTool(BaseTool):
    """Extracts text content from PDF and DOCX documents.

    Validates files before parsing (type, size, magic bytes).
    When sandbox_enabled is True, parsing runs inside Docker for isolation.
    """

    name = "document_parser"
    description = "Parse PDF and DOCX files to extract text"

    def __init__(self, config: dict | None = None) -> None:
        super().__init__(config)
        self.max_file_size_mb = self.config.get("max_file_size_mb", 50)
        self.allowed_types = self.config.get("allowed_file_types", [".pdf", ".docx", ".doc"])
        self.sandbox_enabled = self.config.get("sandbox_enabled", True)

    def execute(self, **kwargs: Any) -> dict:
        """Parse a document and extract text.

        kwargs:
            file_path: Path to the document (required)

        Returns:
            {"text": str, "metadata": dict, "page_count": int, "file_type": str}
            or {"error": str} on failure.
        """
        file_path = kwargs.get("file_path", "")
        if not file_path:
            return {"error": "No file_path provided"}

        path = Path(file_path)

        # Validate before parsing (extension, size, magic bytes)
        is_valid, reason = validate_file(path, self.allowed_types, self.max_file_size_mb)
        if not is_valid:
            logger.warning("document_parser.invalid", path=str(path), reason=reason)
            return {"error": reason}

        suffix = path.suffix.lower()
        file_type = "pdf" if suffix == ".pdf" else "docx"

        try:
            if self.sandbox_enabled:
                return self._parse_in_sandbox(path, file_type)
            else:
                logger.warning("document_parser.sandbox_disabled",
                             msg="Parsing on host — sandbox is disabled")
                if suffix == ".pdf":
                    return self._parse_pdf(path)
                elif suffix in (".docx", ".doc"):
                    return self._parse_docx(path)
                else:
                    return {"error": f"Unsupported file type: {suffix}"}
        except Exception as e:
            logger.error("document_parser.error", path=str(path), error=str(e))
            return {"error": f"Parse failed: {str(e)}"}

    def _parse_in_sandbox(self, path: Path, file_type: str) -> dict:
        """Parse a document inside the Docker sandbox."""
        from secure_agents.core.sandbox import run_in_sandbox

        input_data = {
            "filename": path.name,
            "file_type": file_type,
        }

        # The sandbox script reads from /input/ — we pass the file content
        # via the input data mount. run_in_sandbox handles the mount.
        # But we need to copy the file into the input directory.
        # Actually, run_in_sandbox only supports JSON input_data, not files.
        # We need to extend it or use a different approach.
        # For now: read the file bytes as base64, pass in JSON, decode in sandbox.

        import base64
        file_bytes = path.read_bytes()
        input_data["file_base64"] = base64.b64encode(file_bytes).decode()

        sandbox_script = f"""
import json, base64, sys
from pathlib import Path

input_path = Path("/input/data.json")
data = json.loads(input_path.read_text())
file_type = data["file_type"]
filename = data["filename"]

# Decode file from base64
file_bytes = base64.b64decode(data["file_base64"])
file_path = Path("/tmp") / filename
file_path.write_bytes(file_bytes)

result = {{"text": "", "metadata": {{}}, "page_count": 0, "file_type": file_type}}

try:
    if file_type == "pdf":
        import pdfplumber
        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            result["page_count"] = len(pdf.pages)
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        result["text"] = "\\n\\n".join(text_parts)
    elif file_type in ("docx", "doc"):
        import docx
        doc = docx.Document(str(file_path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        result["text"] = "\\n\\n".join(paragraphs)
        result["page_count"] = max(1, len(result["text"]) // 3000)
    result["metadata"] = {{"filename": filename, "size_bytes": len(file_bytes)}}
except Exception as e:
    result["error"] = str(e)

output_path = Path("/output/result.json")
output_path.write_text(json.dumps(result))
"""
        result = run_in_sandbox(
            script=sandbox_script,
            input_data=input_data,
            timeout=120,
            sandbox_enabled=True,
        )
        if "error" in result:
            logger.error("document_parser.sandbox_error", error=result["error"])
        else:
            logger.info("document_parser.sandbox_ok",
                       path=str(path), chars=len(result.get("text", "")))
        return result

    def _parse_pdf(self, path: Path) -> dict:
        """Extract text from a PDF file using pdfplumber (host-only fallback)."""
        import pdfplumber

        text_parts = []
        page_count = 0

        with pdfplumber.open(path) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

        text = "\n\n".join(text_parts)
        logger.info("document_parser.pdf", path=str(path), pages=page_count, chars=len(text))

        return {
            "text": text,
            "metadata": {"filename": path.name, "size_bytes": path.stat().st_size},
            "page_count": page_count,
            "file_type": "pdf",
        }

    def _parse_docx(self, path: Path) -> dict:
        """Extract text from a DOCX file using python-docx (host-only fallback)."""
        import docx

        doc = docx.Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        text = "\n\n".join(paragraphs)
        page_count = max(1, len(text) // 3000)  # Rough estimate

        logger.info("document_parser.docx", path=str(path), paragraphs=len(paragraphs), chars=len(text))

        return {
            "text": text,
            "metadata": {"filename": path.name, "size_bytes": path.stat().st_size},
            "page_count": page_count,
            "file_type": "docx",
        }

    def validate_config(self) -> bool:
        return True
